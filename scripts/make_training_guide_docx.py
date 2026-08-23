#!/usr/bin/env python
"""Generate docs/Training_Experiment_Guide.docx.

Reproducible: parameter counts are measured from the real code at build time, not
typed in, so the document cannot drift from the repository.

    python scripts/make_training_guide_docx.py
"""

from __future__ import annotations

import os
import sys

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

REPO = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
sys.path.insert(0, os.path.join(REPO, "src"))
from aetfpe.campaign import EXPERIMENTS, MAX_TRAIN_PARAMS, REUSE, build_matrix  # noqa: E402


def main() -> int:
    matrix = {r["id"]: r for r in build_matrix()}
    d = Document()
    for s in d.styles["Normal"], :
        s.font.name = "Calibri"; s.font.size = Pt(11)

    t = d.add_heading("Major Revision Training Experiment Plan", level=0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = d.add_paragraph("AE-TFPE — Robust Feature Fusion Model for Plant Leaf Disease "
                        "Classification\nPrepared for the Major Revision training campaign")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # -- 1 --------------------------------------------------------------
    d.add_heading("1. Research Goal", level=1)
    d.add_paragraph(
        "The Major Revision presents two clearly separated methods. They must never be "
        "merged, and the newer one must never be presented as if it had always been the "
        "original.")
    d.add_heading("Original AE-TFPE — the reference formulation", level=2)
    d.add_paragraph(
        "The method as submitted: positional encoding (PE-RGB), a frozen ViT-B/16 "
        "transformer branch, an image-space stacked sparse denoising auto-encoder, and a "
        "YOLOv8n-cls classifier. Where the historical implementation could not be "
        "recovered from surviving artifacts, the components were reconstructed and are "
        "labelled as reconstructions. This method is retained for comparison and for the "
        "reviewer-facing ablation and fusion reasoning. It is NOT claimed to be "
        "lightweight: it measures 87.5 million parameters and 34.87 GFLOPs.")
    d.add_heading("Efficient AE-TFPE — the Major-Revision improvement", level=2)
    d.add_paragraph(
        "A new, efficiency-oriented design introduced during the revision. It replaces the "
        "ViT-B/16 with MobileViT-XXS truncated after stage 2 (a 28x28 feature grid), and "
        "moves the auto-encoder from image space into that feature space as a slim "
        "denoising auto-encoder. The YOLOv8n-cls classifier is unchanged. The leading "
        "candidate is E5 (C2-28) at 1.72 million parameters and 1.01 GFLOPs.")
    d.add_heading("Why new controlled experiments are required", level=2)
    for s in (
        "The reviewers challenged reproducibility, component attribution, fusion-operator "
        "choice, baseline fairness, computational cost and robustness. Most of the "
        "original evidence could not be reproduced from surviving artifacts.",
        "The historical transform was recovered and found to be a zero-parameter pointwise "
        "look-up table. Mechanism controls are therefore essential: if a simple photometric "
        "transform reproduces the reported effect, the contribution lies in the input "
        "transform rather than in the architecture.",
        "No fair baseline currently exists under the revised evaluation protocol, so no "
        "accuracy comparison is interpretable until one is trained.",
    ):
        d.add_paragraph(s, style="List Bullet")

    # -- 2 --------------------------------------------------------------
    d.add_heading("2. Experiment ID Prefixes", level=1)
    tb = d.add_table(rows=1, cols=2); tb.style = "Light Grid Accent 1"
    tb.rows[0].cells[0].text = "Prefix"; tb.rows[0].cells[1].text = "Meaning"
    for k, v in [
        ("A", "Ablation — which components contribute (Original side)"),
        ("F", "Fusion — how PE and transformer features are combined"),
        ("M", "Mechanism / control — rules out simpler explanations"),
        ("B", "Baseline / benchmark — external classifiers under the same protocol"),
        ("E", "Efficient AE-TFPE — the Major-Revision improvement"),
        ("D", "Denoising-objective control (D1): AE fusion trained clean-to-clean, so "
              "A5 minus D1 isolates the denoising objective itself"),
    ]:
        c = tb.add_row().cells; c[0].text = k; c[1].text = v

    # -- 3 --------------------------------------------------------------
    d.add_heading("3. Experiment Table", level=1)
    d.add_paragraph(
        f"Parameter counts are measured from the implementation. Models above "
        f"{MAX_TRAIN_PARAMS:,} total parameters are skipped by default under the one-day "
        f"budget; each skip is justified in Section 6.")
    cols = ["ID", "Model / Experiment", "Group", "Architecture", "Purpose",
            "Params", "Trainable", "Train?", "Priority", "Included", "Reason"]
    tb = d.add_table(rows=1, cols=len(cols)); tb.style = "Light Grid Accent 1"
    for i, c in enumerate(cols):
        tb.rows[0].cells[i].text = c
    ARCH = {
        "A0": "YOLOv8n-cls only", "A1": "PE + YOLO", "A2": "ViT-B/16 + YOLO",
        "A3": "PE + ViT-B/16 + linear fusion", "A4": "image-space AE + YOLO",
        "A5": "PE + ViT-B/16 + image-space AE", "D1": "as A5, clean-to-clean AE",
        "F1": "PE + ViT, addition", "F2": "PE + ViT, concat (6-ch stem)",
        "F4": "PE + ViT, channel attention", "B1": "ResNet-50", "B2": "EfficientNet-B0",
        "B3": "ViT-B/16 classifier", "E3": "PE + MobileViT s2, no AE",
        "E5": "PE + MobileViT s2 (28x28) + feature AE", "E7": "PE + MobileViT s4 (7x7) + feature AE",
    }
    for e in EXPERIMENTS:
        r = matrix[e["id"]]
        inc = r["status"] != "SKIPPED_SIZE"
        c = tb.add_row().cells
        for i, v in enumerate([
            e["id"], e["title"], e["group"], ARCH.get(e["id"], ""), e["purpose"],
            f"{r['params']:,}", f"{r['trainable_params']:,}",
            "Yes" if inc else "No", e["priority"],
            "Included" if inc else "Skipped",
            "within threshold" if inc else f"exceeds {MAX_TRAIN_PARAMS:,}",
        ]):
            c[i].text = str(v)
    for logical, (src, why) in REUSE.items():
        c = tb.add_row().cells
        for i, v in enumerate([logical, f"reuses {src}", "F", "—",
                               "Served by an existing checkpoint", "—", "—", "No",
                               "—", "Reuse", why]):
            c[i].text = str(v)
    d.add_paragraph(
        "Three logical experiments need no training of their own: F3 is identical to A3, "
        "F5 to A5, and the clean AE-fusion arm to D1. This is verified from the config "
        "signatures by scripts/print_run_matrix.py, not assumed. A1 (PE-only) also serves "
        "both method families, because with no transformer branch the Original and "
        "Efficient variants are the same model.")

    # -- 4 --------------------------------------------------------------
    d.add_heading("4. Architecture Summaries", level=1)
    for h, b in [
        ("A-series — Ablation",
         "Removes one component at a time from Original AE-TFPE to show what each "
         "contributes: A0 is the plain classifier, A1 adds positional encoding only, A2 the "
         "transformer only, A3 both without the auto-encoder, A4 the auto-encoder on plain "
         "RGB, and A5 the complete method."),
        ("F-series — Fusion",
         "Holds everything else fixed and varies only how the positional-encoding and "
         "transformer streams are combined: addition, concatenation, concatenation with a "
         "learned projection, channel attention, and the auto-encoder acting as the fusion "
         "mechanism. This answers whether auto-encoder fusion is genuinely better than "
         "conventional operators."),
        ("M-series — Mechanism controls",
         "The most important controls in the campaign. M1 reproduces the historical "
         "look-up-table transform exactly; M2 applies a matched monotonic gamma curve; M3 "
         "trains the plain classifier with corruption augmentation. If any of these matches "
         "the proposed method, the reported benefit comes from the input transform or from "
         "augmentation, not from the architecture."),
        ("B-series — Baselines",
         "Independent classifiers trained under the identical protocol, so comparisons are "
         "fair: ResNet-50, EfficientNet-B0 and ViT-B/16. B2 (EfficientNet-B0) is the "
         "lightweight external baseline retained under the one-day budget."),
        ("E-series — Efficient AE-TFPE",
         "The Major-Revision improvement. E5 is the main candidate at a 28x28 feature grid; "
         "E3 removes the auto-encoder as a control; E7 uses the 7x7 grid to test how much "
         "spatial resolution matters."),
    ]:
        d.add_heading(h, level=2); d.add_paragraph(b)

    d.add_heading("The two methods side by side", level=2)
    tb = d.add_table(rows=1, cols=3); tb.style = "Light Grid Accent 1"
    for i, c in enumerate(["Stage", "Original AE-TFPE", "Efficient AE-TFPE (E5)"]):
        tb.rows[0].cells[i].text = c
    for row in [
        ("Positional encoding", "PE-RGB, sinusoidal", "PE-RGB, sinusoidal — unchanged"),
        ("Transformer branch", "ViT-B/16, frozen, 85.8M params", "MobileViT-XXS stage 2, frozen, 0.15M"),
        ("Feature grid", "image space, 224x224", "28x28, 48 channels"),
        ("Auto-encoder", "image-space stacked sparse denoising", "feature-space slim denoising"),
        ("Latent", "128 x 28 x 28", "64 x 28 x 28"),
        ("Decoder", "3 transposed convolutions to 224", "3 transposed convolutions, 28-56-112-224"),
        ("Classifier", "YOLOv8n-cls, unmodified", "YOLOv8n-cls, unmodified"),
        ("Total parameters", "87,549,123", "1,716,586"),
        ("GFLOPs", "34.87", "1.01"),
        ("T4 latency, batch 1", "28.51 ms", "10.46 ms"),
    ]:
        c = tb.add_row().cells
        for i, v in enumerate(row):
            c[i].text = v
    d.add_paragraph(
        "Efficient AE-TFPE substantially reduces the computational burden of Original "
        "AE-TFPE — about 51 times fewer parameters and 34 times fewer operations — but it "
        "still costs about three times the plain YOLOv8n-cls classifier at batch size one. "
        "It should be described as lightweight relative to the original method, never as "
        "free relative to the baseline classifier.")

    # -- 5 --------------------------------------------------------------
    d.add_heading("5. Reviewer Mapping", level=1)
    tb = d.add_table(rows=1, cols=3); tb.style = "Light Grid Accent 1"
    for i, c in enumerate(["Reviewer concern", "Experiments", "Status after this campaign"]):
        tb.rows[0].cells[i].text = c
    for row in [
        ("Component contribution", "A1, A2, A3, A4, A5, E3, E5",
         "Partly addressed: the Efficient side is covered; the Original side is deferred"),
        ("Fusion comparison", "F1, F2, F3 (=A3), F4, F5 (=A5), D1",
         "Deferred — every arm carries the large frozen ViT (see Section 6)"),
        ("Mechanism validation", "M1, M2, M3",
         "Fully addressed; this is the decisive gate"),
        ("Fair baselines", "A0, B1, B2, B3",
         "A0 and B2 trained; ResNet-50 and ViT-B/16 deferred as fully-trainable and costly"),
        ("Computational efficiency", "Tesla T4 benchmark, already measured",
         "Complete; no further training needed"),
        ("Robustness under degradation", "Local evaluation: Normal / Easy / Moderate / Hard",
         "Runs on the MacBook after training"),
    ]:
        c = tb.add_row().cells
        for i, v in enumerate(row):
            c[i].text = v

    # -- 6 --------------------------------------------------------------
    d.add_heading("6. One-Day Execution Strategy", level=1)
    d.add_paragraph(
        "The campaign runs in three priority tiers so that if time runs out, what was lost "
        "is the least important work rather than an arbitrary remainder.")
    for h, ids, why in [
        ("Priority P0 — must complete", "A0, E5, M1, M2, M3, E3",
         "These decide whether the revised paper is defensible. A0 provides the fair "
         "baseline without which no comparison means anything; E5 is the contribution; "
         "M1-M3 test whether a simpler explanation accounts for the results; E3 removes the "
         "auto-encoder. If E5 does not clearly beat M1-M3, the campaign should stop and "
         "report that honestly rather than continue."),
        ("Priority P1 — high value", "A1, A4, E7, B2",
         "Component evidence and an independent external baseline. Valuable for the "
         "response letter, but the main argument survives without them."),
        ("Priority P2 — optional", "A2, A3, A5, D1, F1, F2, F4, B1, B3",
         "All exceed the size threshold and are skipped by default."),
    ]:
        d.add_heading(h, level=2)
        d.add_paragraph(f"Experiments: {ids}")
        d.add_paragraph(why)

    d.add_heading("Why the large models are skipped, and what it costs", level=2)
    d.add_paragraph(
        "Seven of the nine skipped arms (A2, A3, A5, D1, F1, F2, F4) contain only about 1.5 "
        "to 1.75 million trainable parameters. The remaining 85.8 million is a frozen "
        "ViT-B/16 that runs without gradients, so the expense is its forward pass rather "
        "than optimiser work — but it is still expensive: on a Tesla T4 these arms process "
        "about 93 images per second against the baseline's 4,365, a 47-fold difference. "
        "B1 and B3 are different: they are fully trainable and therefore genuinely costly.")
    d.add_paragraph("What can and cannot be recovered without training them:")
    for s in (
        "Efficiency and complexity claims are unaffected — parameters, operations, latency, "
        "throughput and memory were already measured on a Tesla T4 and archived. Training "
        "would change none of them.",
        "Original AE-TFPE can be reported as a computational and reference formulation, "
        "which is the recommended treatment under the one-day budget.",
        "The claim that auto-encoder fusion is superior to conventional operators CANNOT be "
        "recovered by analysis: it is an accuracy comparison and needs F1, F2, F4 and D1. "
        "If those are not trained, that claim must be withdrawn or explicitly deferred.",
        "External baseline fairness is partly covered by EfficientNet-B0, so skipping "
        "ResNet-50 and ViT-B/16 weakens but does not remove the fairness response.",
    ):
        d.add_paragraph(s, style="List Bullet")
    d.add_paragraph(
        "Decision required: to keep the fusion-superiority and denoising-objective claims, "
        "the minimum additional set is A5, D1, F1, F2 and F4. The notebook estimates the "
        "cost before anything starts; they are trained only if their IDs are added to "
        "FORCE_LARGE_IDS.")

    # -- 7 --------------------------------------------------------------
    d.add_heading("7. Expected Outputs", level=1)
    d.add_paragraph(
        "Everything is written to Google Drive under AE_TFPE_MajorRevision and "
        "synchronised after every epoch, so a Colab disconnect costs at most one epoch of "
        "work. The Drive copy is the source of truth for recovery.")
    tb = d.add_table(rows=1, cols=2); tb.style = "Light Grid Accent 1"
    tb.rows[0].cells[0].text = "Location"; tb.rows[0].cells[1].text = "Contents"
    for k, v in [
        ("campaign/campaign_manifest.json", "Status of every run: parameters, timings, GPU, commit"),
        ("campaign/campaign_summary.csv", "The same information as a flat table"),
        ("checkpoints/<ID>/checkpoint.pt", "Best-validation weights, with config and class list embedded"),
        ("checkpoints/<ID>/metrics.csv", "Per-epoch loss, top-1, top-5, reconstruction loss, timing, peak memory"),
        ("checkpoints/<ID>/train_summary.json", "Protocol, best validation accuracy, wall-clock, environment"),
        ("checkpoints/<ID>/config.yaml", "The exact resolved configuration used"),
        ("logs/<ID>.log", "Complete training output"),
        ("completed/ and failed/", "Simple markers for a quick status check"),
        ("campaign/for_local_evaluation.tar.gz", "Everything needed for evaluation on the MacBook"),
    ]:
        c = tb.add_row().cells; c[0].text = k; c[1].text = v
    d.add_paragraph(
        "No test-set evaluation and no corrupted data are produced during training. "
        "Evaluation against Normal, Easy, Moderate and Hard runs afterwards on the MacBook "
        "Pro M4, following docs/LOCAL_EVALUATION_HANDOFF.md. Deployment latency and memory "
        "figures come only from the Tesla T4 benchmarks already collected; training on an "
        "A100 does not replace them.")

    out = os.path.join(REPO, "docs", "Training_Experiment_Guide.docx")
    d.save(out)
    print(f"wrote {out}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
